#!/usr/bin/env python3
"""
Transcript Processor
Parses a production rundown .docx file, translates content, synthesizes TTS,
and assembles a time-aligned WAV file per language.

Requires: sudo apt install ffmpeg (for pydub WAV export)

Usage:
    python transcript_processor.py rundown.docx
    python transcript_processor.py rundown.docx --output ./output --languages cs sv
"""

import argparse
import io
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from pydub import AudioSegment
from pydub.effects import speedup
from dotenv import load_dotenv

from google.cloud import texttospeech
from google.cloud import translate_v2 as translate

load_dotenv()

# ── Language configuration ────────────────────────────────────────────────────

LANGUAGES = {
    'cs': {
        'name': 'Czech',
        'translate_code': 'cs',
        'tts_language': 'cs-CZ',
        'voice_female': 'cs-CZ-Wavenet-A',
        'voice_male': 'cs-CZ-Wavenet-A',   # no male Czech Wavenet voice
    },
    'sv': {
        'name': 'Swedish',
        'translate_code': 'sv',
        'tts_language': 'sv-SE',
        'voice_female': 'sv-SE-Wavenet-A',
        'voice_male': 'sv-SE-Wavenet-B',
    },
}

# Extend these lists as new speakers appear in future documents
MALE_NAMES = {'dag', 'orjan', 'ørjan', 'ivan', 'jonas', 'erik', 'lars', 'magnus', 'bjorn', 'bjørn'}
FEMALE_NAMES = {'hilde', 'maya', 'camilla', 'daniela', 'anna', 'maria', 'sara', 'lisa', 'ingrid'}

TRANSLATABLE_TYPES = {'video', 'vignette'}

# Detects "Name: text" or "Name : text" at paragraph start
NAME_COLON_RE = re.compile(
    r'^([A-ZÆØÅÄÖÜ][a-zA-ZÆØÅÄÖÜæøåäöü\-]{1,20}(?:\s[A-ZÆØÅÄÖÜ][a-zA-ZÆØÅÄÖÜæøåäöü\-]{1,20}){0,2})\s*:\s*(.+)$',
    re.DOTALL,
)

TBL_TAG = qn('w:tbl')
P_TAG = qn('w:p')

# ── Timestamp / duration helpers ──────────────────────────────────────────────

def parse_timestamp(ts_str: str) -> float:
    """Parse HH:MM:SS or HH:MM into seconds since midnight. Tolerates spaces."""
    ts_str = ts_str.strip().replace(' ', '')
    parts = ts_str.split(':')
    if len(parts) == 3:
        h, m, s = int(parts[0]), int(parts[1]), int(parts[2])
    elif len(parts) == 2:
        h, m, s = int(parts[0]), int(parts[1]), 0
    else:
        raise ValueError(f"Cannot parse timestamp: {ts_str!r}")
    return h * 3600 + m * 60 + s


def parse_duration(dur_str: str) -> int:
    """Parse '1m', '30s', '1m30', '1m30s' etc. into seconds."""
    dur_str = dur_str.strip().lower().rstrip('s')  # strip trailing 's' if present
    if 'm' in dur_str:
        parts = dur_str.split('m')
        minutes = int(parts[0])
        seconds = int(parts[1]) if parts[1] else 0
        return minutes * 60 + seconds
    if dur_str.isdigit():
        return int(dur_str)
    return 0

# ── Paragraph analysis ────────────────────────────────────────────────────────

def paragraph_is_bold(para: DocxParagraph) -> bool:
    """True if every non-empty run in the paragraph is bold."""
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


# Matches any content inside (...) or [...] — always stage directions in rundowns
_STAGE_DIR_RE = re.compile(r'\[.*?\]|\(.*?\)', re.DOTALL)

def strip_stage_directions(text: str) -> str:
    """Remove all (...) and [...] blocks from text, collapse extra whitespace."""
    cleaned = _STAGE_DIR_RE.sub('', text)
    return re.sub(r'\s+', ' ', cleaned).strip()


def get_voice_gender(name: str) -> str:
    name_lower = name.lower().strip()
    if name_lower in MALE_NAMES:
        return 'male'
    if name_lower in FEMALE_NAMES:
        return 'female'
    return 'female'   # default

# ── Rundown table parsing ─────────────────────────────────────────────────────

def extract_type(item_str: str) -> str:
    """
    Extract the canonical type keyword from an Item cell.
    'VIDEO - VB3', 'VIDEO– VB4', 'VIDEO- VB 3 Intro' → 'video'
    'Stinger' → 'stinger', 'VIGNETTE' → 'vignette', 'Bumper' → 'bumper'
    """
    s = item_str.lower().strip()
    for keyword in ('video', 'vignette', 'stinger', 'bumper'):
        if s.startswith(keyword):
            return keyword
    return s  # fallback: return full string


def parse_row_cells(cells: list[str]) -> dict | None:
    """
    Parse one rundown row from a list of cell text values.
    Expected: [number, item/type, timestamp, duration, title, ...]
    Returns None if cells don't look like a valid rundown row.
    """
    if len(cells) < 5:
        return None
    if not cells[0].strip().isdigit():
        return None
    try:
        ts = parse_timestamp(cells[2])
    except (ValueError, IndexError):
        return None
    return {
        'number': int(cells[0].strip()),
        'type': extract_type(cells[1]),
        'timestamp_s': ts,
        'duration_s': parse_duration(cells[3]),
        'title': cells[4].split('\n')[0].strip(),  # first line only as title
        'utterances': [],
        'slot_ms': 0,
    }


def parse_table_rows(table: DocxTable) -> list[dict]:
    """
    Parse all valid rundown rows from a table (handles single and multi-row tables).
    For multi-row tables, content in extra cell paragraphs is also extracted.
    """
    results = []
    for row in table.rows:
        cells_text = [c.text.strip() for c in row.cells]
        seg = parse_row_cells(cells_text)
        if not seg:
            continue
        # For multi-row tables: content may live in extra paragraphs of the last cell
        last_cell = row.cells[-1]
        if len(last_cell.paragraphs) > 1:
            seg['utterances'] = parse_content_paragraphs(last_cell.paragraphs[1:])
        results.append(seg)
    return results

# ── Content paragraph parsing ─────────────────────────────────────────────────

def parse_content_paragraphs(paragraphs: list[DocxParagraph]) -> list[dict]:
    """
    Convert content paragraphs into a list of utterances.
    Each utterance: {'text': str, 'gender': str}
    """
    utterances = []
    current_gender = 'female'

    for para in paragraphs:
        raw = para.text.strip()
        if not raw:
            continue
        # Bold paragraph → skip entirely (clarification title or speaker name label)
        if paragraph_is_bold(para):
            words = raw.split()
            if 1 <= len(words) <= 3 and all(w[0].isupper() for w in words if w):
                current_gender = get_voice_gender(words[0])
            continue
        # Strip (...) and [...] stage directions from the line
        text = strip_stage_directions(raw)
        if not text:
            continue
        # "Name:" alone — stage direction stripped everything after colon → voice switch only
        if re.match(r'^[A-ZÆØÅÄÖÜ][a-zA-ZÆØÅÄÖÜæøåäöü]{1,20}(?:\s[A-ZÆØÅÄÖÜ][a-zA-ZÆØÅÄÖÜæøåäöü]{1,20}){0,2}:$', text):
            current_gender = get_voice_gender(text[:-1].strip())
            continue
        # "Name: text" pattern → voice switch, translate only the text part
        # Also handles "Name (stage dir): text" after stripping, e.g. "Ivan (off): Hello"
        match = NAME_COLON_RE.match(text)
        if match:
            name = match.group(1).strip()
            content = match.group(2).strip()
            current_gender = get_voice_gender(name)
            if content:
                utterances.append({'text': content, 'gender': current_gender})
            continue
        # Plain text → translate
        utterances.append({'text': text, 'gender': current_gender})

    return utterances

# ── Full document parse ───────────────────────────────────────────────────────

def parse_rundown(docx_path: Path, debug: bool = False) -> list[dict]:
    """
    Parse the rundown .docx into an ordered list of segments.
    Each segment has type, timestamp, slot_ms, utterances.
    """
    doc = Document(docx_path)
    segments: list[dict] = []
    current_seg: dict | None = None
    pending_paras: list[DocxParagraph] = []

    def flush():
        if current_seg is not None and pending_paras:
            parsed = parse_content_paragraphs(pending_paras)
            # Only overwrite if we found something (cell-level content takes precedence)
            if parsed:
                current_seg['utterances'] = parsed

    for child in doc.element.body:
        if child.tag == TBL_TAG:
            flush()
            pending_paras = []
            table = DocxTable(child, doc)
            rows = parse_table_rows(table)
            if debug:
                print(f"  [TABLE] {len(table.rows)} rows → {len(rows)} valid segments parsed")
                for r in rows:
                    print(f"    row {r['number']}: type={r['type']!r}  ts={r['timestamp_s']}s  utterances={len(r['utterances'])}")
            if rows:
                segments.extend(rows)
                current_seg = rows[-1]
        elif child.tag == P_TAG:
            para = DocxParagraph(child, doc)
            text = para.text.strip()
            if debug and text:
                bold = paragraph_is_bold(para)
                print(f"  [PARA ] bold={bold}  {text[:70]!r}")
            if current_seg is not None:
                pending_paras.append(para)

    flush()

    # Calculate slot_ms from timestamp deltas
    for i, seg in enumerate(segments):
        if i + 1 < len(segments):
            delta_s = segments[i + 1]['timestamp_s'] - seg['timestamp_s']
            seg['slot_ms'] = max(0, int(delta_s * 1000))
        else:
            seg['slot_ms'] = seg['duration_s'] * 1000

    return segments

# ── Translation & TTS ─────────────────────────────────────────────────────────

def translate_text(text: str, target_lang: str, translate_client) -> str:
    result = translate_client.translate(text, target_language=target_lang)
    return result['translatedText']


def synthesize_clip(text: str, lang_cfg: dict, gender: str, tts_client) -> AudioSegment:
    voice_name = lang_cfg['voice_male'] if gender == 'male' else lang_cfg['voice_female']
    response = tts_client.synthesize_speech(
        input=texttospeech.SynthesisInput(text=text),
        voice=texttospeech.VoiceSelectionParams(
            language_code=lang_cfg['tts_language'],
            name=voice_name,
        ),
        audio_config=texttospeech.AudioConfig(
            audio_encoding=texttospeech.AudioEncoding.MP3,
            speaking_rate=1.0,
        ),
    )
    return AudioSegment.from_mp3(io.BytesIO(response.audio_content))


def fit_to_slot(audio: AudioSegment, slot_ms: int) -> AudioSegment:
    """Speed up audio only if it exceeds the slot. Pad with silence if shorter."""
    if len(audio) > slot_ms:
        rate = len(audio) / slot_ms
        audio = speedup(audio, playback_speed=rate, chunk_size=150, crossfade=25)
    silence_needed = slot_ms - len(audio)
    if silence_needed > 0:
        audio = audio + AudioSegment.silent(duration=silence_needed)
    return audio

# ── Audio assembly ────────────────────────────────────────────────────────────

def trim_to_first_translatable(segments: list[dict]) -> list[dict]:
    """Drop all leading segments before the first translatable one with utterances."""
    for i, seg in enumerate(segments):
        if seg['type'] in TRANSLATABLE_TYPES and seg.get('utterances'):
            return segments[i:]
    return segments


def build_language_audio(
    segments: list[dict],
    lang_key: str,
    translate_client,
    tts_client,
) -> AudioSegment:
    lang_cfg = LANGUAGES[lang_key]
    timeline = AudioSegment.empty()

    for seg in segments:
        slot_ms = seg['slot_ms']
        if slot_ms == 0:
            continue

        seg_type = seg['type']
        utterances = seg.get('utterances', [])

        if seg_type not in TRANSLATABLE_TYPES or not utterances:
            # Stinger, bumper, or empty translatable row → silence
            timeline += AudioSegment.silent(duration=slot_ms)
            continue

        print(f"  [{seg['number']:>3}] {seg['title'][:55]}")

        seg_audio = AudioSegment.empty()
        for utt in utterances:
            try:
                translated = translate_text(
                    utt['text'], lang_cfg['translate_code'], translate_client
                )
                clip = synthesize_clip(translated, lang_cfg, utt['gender'], tts_client)
                seg_audio += clip
            except Exception as exc:
                print(f"         ⚠ Skipped utterance: {exc}")

        if len(seg_audio) == 0:
            timeline += AudioSegment.silent(duration=slot_ms)
        else:
            timeline += fit_to_slot(seg_audio, slot_ms)

    return timeline

# ── Dry run ───────────────────────────────────────────────────────────────────

def dry_run(segments: list[dict]) -> None:
    """Print a summary of what would be translated without calling any APIs."""
    print("\n── DRY RUN — no API calls made " + "─" * 35)
    translatable_count = 0
    utterance_count = 0

    for seg in segments:
        slot_s = seg['slot_ms'] / 1000
        seg_type = seg['type']
        utterances = seg.get('utterances', [])

        if seg_type not in TRANSLATABLE_TYPES:
            print(f"  [{seg['number']:>3}] SKIP  {seg_type:<10}  slot={slot_s:.0f}s")
            continue

        if not utterances:
            print(f"  [{seg['number']:>3}] EMPTY {seg_type:<10}  slot={slot_s:.0f}s  \"{seg['title'][:40]}\"")
            continue

        translatable_count += 1
        print(f"\n  [{seg['number']:>3}] {seg_type.upper():<10}  slot={slot_s:.0f}s  \"{seg['title'][:40]}\"")
        for utt in utterances:
            utterance_count += 1
            speaker = f"[{utt['gender']}]"
            preview = utt['text'][:80].replace('\n', ' ')
            ellipsis = '…' if len(utt['text']) > 80 else ''
            print(f"         {speaker}  {preview}{ellipsis}")

    print(f"\n── Summary " + "─" * 50)
    print(f"  Translatable segments : {translatable_count}")
    print(f"  Utterances to process : {utterance_count}")
    total_s = sum(s['slot_ms'] for s in segments) / 1000
    print(f"  Total timeline        : {total_s/60:.1f} minutes")
    print()

# ── CLI entry point ───────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='Translate and synthesize a rundown .docx transcript.'
    )
    parser.add_argument('docx', help='Path to the rundown .docx file')
    parser.add_argument('--output', default='./output', help='Output directory (default: ./output)')
    parser.add_argument(
        '--languages', nargs='+', default=['cs', 'sv'],
        choices=list(LANGUAGES.keys()),
        help='Target languages (default: cs sv)',
    )
    parser.add_argument(
        '--dry-run', action='store_true',
        help='Parse and preview what would be translated — no API calls',
    )
    parser.add_argument(
        '--debug', action='store_true',
        help='Print every table and paragraph found during parsing',
    )
    parser.add_argument(
        '--trim-start', action='store_true',
        help='Skip silence before the first translatable segment (removes pre-show slots)',
    )
    args = parser.parse_args()

    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"Error: file not found: {docx_path}")
        sys.exit(1)

    print(f"\nParsing: {docx_path.name}")
    if args.debug:
        print("── DEBUG output ──────────────────────────────────────────────")
    segments = parse_rundown(docx_path, debug=args.debug)
    if args.debug:
        print("─────────────────────────────────────────────────────────────\n")
    if args.trim_start:
        segments = trim_to_first_translatable(segments)

    if args.dry_run:
        dry_run(segments)
        return

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    translatable = [s for s in segments if s['type'] in TRANSLATABLE_TYPES and s['utterances']]
    total_s = sum(s['slot_ms'] for s in segments) / 1000
    print(f"Rows: {len(segments)} total, {len(translatable)} with translatable content")
    print(f"Total timeline: {total_s/60:.1f} minutes\n")

    translate_client = translate.Client()
    tts_client = texttospeech.TextToSpeechClient()

    for lang_key in args.languages:
        lang_name = LANGUAGES[lang_key]['name']
        print(f"── {lang_name} ({'─' * 50})")
        audio = build_language_audio(segments, lang_key, translate_client, tts_client)
        out_path = output_dir / f"{docx_path.stem}_{lang_key}.wav"
        audio.export(str(out_path), format='wav')
        print(f"   → {out_path}  ({len(audio)/1000:.1f}s)\n")

    print("Done.")


if __name__ == '__main__':
    main()
